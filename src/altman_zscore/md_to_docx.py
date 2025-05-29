"""
Markdown to DOCX conversion utility for Altman Z-Score reports.

This module provides the function `convert_report_md_to_docx`, which converts a Markdown report to DOCX format, supporting tables, titles, images, and chart embedding. This enables high-quality, portable reporting for financial analysis workflows.

Usage:
    from altman_zscore.md_to_docx import convert_report_md_to_docx
    convert_report_md_to_docx('path/to/report.md')

Features:
- Converts Markdown to DOCX, preserving headings, tables, lists, and blockquotes
- Embeds images (including Z-Score/price trend charts) if referenced in the Markdown
- Handles all standard Markdown features used in Altman Z-Score reports
- Output DOCX is suitable for sharing, printing, or further editing
"""
import os
from docx import Document
from docx.shared import Inches
import markdown
from bs4 import BeautifulSoup, Tag
from bs4.element import NavigableString

from altman_zscore.reporting import print_info

def convert_report_md_to_docx(md_path, docx_path=None):
    """
    Convert a Markdown report to DOCX format, supporting tables, titles, images, and embedding the chart.
    Args:
        md_path (str): Path to the Markdown file.
        docx_path (str): Output path for the DOCX file. If None, replaces .md with .docx.
    """
    if docx_path is None:
        docx_path = os.path.splitext(md_path)[0] + '.docx'

    with open(md_path, 'r', encoding='utf-8') as f:
        md_content = f.read()

    html = markdown.markdown(md_content, extensions=['tables', 'fenced_code', 'attr_list'])
    soup = BeautifulSoup(html, 'html.parser')
    doc = Document()

    def add_paragraph_with_style(text, style=None):
        if style:
            doc.add_paragraph(text, style=style)
        else:
            doc.add_paragraph(text)

    def handle_table(table_tag):
        rows = table_tag.find_all('tr')
        if not rows:
            return
        cols = rows[0].find_all(['td', 'th'])
        table = doc.add_table(rows=len(rows), cols=len(cols))
        for i, row in enumerate(rows):
            cells = row.find_all(['td', 'th'])
            for j, cell in enumerate(cells):
                table.cell(i, j).text = cell.get_text()

    def handle_img(img_tag):
        img_src = img_tag.get('src', None)
        if not img_src:
            return
        # Try resolving relative to the .md file
        img_path_md = os.path.join(os.path.dirname(md_path), img_src)
        # Try resolving relative to the workspace root (project root)
        workspace_root = os.path.abspath(os.path.join(os.path.dirname(md_path), '../../..'))
        img_path_root = os.path.join(workspace_root, img_src)
        # Try both, prefer the one that exists
        if os.path.exists(img_path_md):
            doc.add_picture(img_path_md, width=Inches(5.5))
        elif os.path.exists(img_path_root):
            doc.add_picture(img_path_root, width=Inches(5.5))
        else:
            add_paragraph_with_style(f'[Image not found: {img_src}]')

    for elem in soup.contents:
        if isinstance(elem, NavigableString):
            if elem.strip():
                add_paragraph_with_style(str(elem))
        elif isinstance(elem, Tag):
            tag = elem.name
            if tag == 'h1':
                add_paragraph_with_style(elem.get_text(), style='Title')
            elif tag == 'h2':
                add_paragraph_with_style(elem.get_text(), style='Heading 1')
            elif tag == 'h3':
                add_paragraph_with_style(elem.get_text(), style='Heading 2')
            elif tag == 'h4':
                add_paragraph_with_style(elem.get_text(), style='Heading 3')
            elif tag == 'p':
                add_paragraph_with_style(elem.get_text())
            elif tag == 'ul':
                for li in elem.find_all('li', recursive=False):
                    add_paragraph_with_style('• ' + li.get_text())
            elif tag == 'ol':
                for idx, li in enumerate(elem.find_all('li', recursive=False), 1):
                    add_paragraph_with_style(f'{idx}. ' + li.get_text())
            elif tag == 'table':
                handle_table(elem)
            elif tag == 'blockquote':
                add_paragraph_with_style(elem.get_text(), style='Intense Quote')
            elif tag == 'img':
                handle_img(elem)
            elif tag == 'hr':
                doc.add_paragraph('---')
            elif tag == 'pre' or tag == 'code':
                add_paragraph_with_style(elem.get_text(), style='Intense Quote')
            # Add more handlers as needed

    doc.save(docx_path)
    print_info(f"DOCX report saved to {docx_path}")
